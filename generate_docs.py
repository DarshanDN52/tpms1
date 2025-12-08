import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_documentation():
    doc = Document()
    
    # --- Title Page ---
    title = doc.add_heading('PCAN & TPMS Project Detailed Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n')
    
    # --- 1. Executive Summary ---
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        'This report details the operational logic, data flow, and architecture of the '
        'Vehicle Health Monitoring System. The system integrates hardware interfacing (PCAN-USB), '
        'wireless sensor protocols (BLE), and real-time visualization (React Frontend) to provide '
        'a complete diagnostic dashboard.'
    )

    # --- 2. System Architecture ---
    doc.add_heading('2. System Architecture', level=1)
    doc.add_paragraph(
        'The application is built on a decoupled Client-Server architecture.'
    )
    
    # Backend Section
    doc.add_heading('2.1. Backend Layer (Python/FastAPI)', level=2)
    p = doc.add_paragraph()
    p.add_run('Core Responsibility: ').bold = True
    p.add_run('Hardware abstraction and API serving.')
    
    doc.add_paragraph('Key Components:', style='List Bullet')
    doc.add_paragraph('FastAPI Server: Asynchronous web server handling HTTP requests and WebSockets.', style='List Bullet 2')
    doc.add_paragraph('PCAN Service: A singleton service managing the dedicated thread for reading CAN bus messages via the PCANBasic API.', style='List Bullet 2')
    doc.add_paragraph('BLE Automation Module: A task-based system that manages long-running Bluetooth Low Energy connection sequences and testing.', style='List Bullet 2')

    # Frontend Section
    doc.add_heading('2.2. Frontend Layer (React/Vite)', level=3)
    p = doc.add_paragraph()
    p.add_run('Core Responsibility: ').bold = True
    p.add_run('User interaction, state management, and real-time rendering.')
    
    doc.add_paragraph('Key Components:', style='List Bullet')
    doc.add_paragraph('Dashboard Page: Renders the 3D Truck model and charts using Chart.js.', style='List Bullet 2')
    doc.add_paragraph('CAN Console: Provides a low-level debugging interface for raw hex messages.', style='List Bullet 2')
    doc.add_paragraph('BLE Test Control: Interface for configuring and running automated test suites.', style='List Bullet 2')

    # --- 3. Operational Data Flows ---
    doc.add_heading('3. Operational Data Flows', level=1)
    
    # 3.1 Startup Flow
    doc.add_heading('3.1. System Startup & Initialization', level=2)
    doc.add_paragraph(
        '1. The Backend starts and initializes the `PCANService` and `TPMSService` singletons.\n'
        '2. The Frontend loads and checks the API health status (`/api/pcan/status`).\n'
        '3. If PCAN hardware is detected, the Frontend automatically requests initialization with default settings (500k baud).\n'
        '4. A WebSocket connection is established at `ws://<host>/ws` to listen for global broadcast events (e.g., BLE logs).'
    )

    # 3.2 CAN Data Flow
    doc.add_heading('3.2. CAN Bus Data Pipeline', level=2)
    doc.add_paragraph(
        'This flow describes how data moves from the vehicle bus to the user screen:'
    )
    doc.add_paragraph('Step 1: Hardware Interrupt', style='List Number')
    doc.add_paragraph('The PCAN-USB driver receives a CAN frame from the physical bus.', style='List Paragraph')
    
    doc.add_paragraph('Step 2: Service Polling', style='List Number')
    doc.add_paragraph('The `PCANService` background thread runs a continuous loop (microsleep 0.05s) calling `Read()`. Received messages are pushed into a thread-safe `deque` buffer.', style='List Paragraph')
    
    doc.add_paragraph('Step 3: Frontend Fetch', style='List Number')
    doc.add_paragraph('The React Frontend polls the `/api/pcan/read` endpoint every 50ms.', style='List Paragraph')
    
    doc.add_paragraph('Step 4: Visualization', style='List Number')
    doc.add_paragraph('New messages are processed by `TPMSDashboard.jsx`. If the message ID matches a configured Tire ID, the payload is parsed for Pressure/Temperature/Battery and the state is updated.', style='List Paragraph')

    # 3.3 BLE Test Flow
    doc.add_heading('3.3. BLE Automation Flow', level=2)
    doc.add_paragraph(
        '1. User selects a target device and test configuration (UUIDs, retries) in the UI.\n'
        '2. A POST request to `/start-test` initiates an asynchronous task in `main.py`.\n'
        '3. The `BLETestAutomation` class connects to the device using `BleakClient`.\n'
        '4. Commands are executed in chunks. For each command:\n'
        '    a. Command is written to the WRITE characteristic.\n'
        '    b. System waits for a Notification response.\n'
        '    c. Response is validated (Pass/Fail) and logged.\n'
        '5. Every log entry is broadcast immediately via WebSocket to the frontend log window.'
    )

    # --- 4. Logic & Threshold Specifications ---
    doc.add_heading('4. Detailed Logic Specifications', level=1)
    
    doc.add_heading('4.1. Tire Status Classification', level=2)
    doc.add_paragraph('The system visually indicates tire health using color codes derived from these specific thresholds:')
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Metric'
    hdr[1].text = 'Critical (Red)\nAction Required'
    hdr[2].text = 'Warning (Orange)\nCaution'
    hdr[3].text = 'Normal (Green)\nGood'
    
    def add_row(metric, crit, warn, norm):
        row = table.add_row().cells
        row[0].text = metric
        row[1].text = crit
        row[2].text = warn
        row[3].text = norm

    add_row('Pressure (PSI)', '< 20  OR  > 120', '< 30  OR  > 100', '30 - 100')
    add_row('Temperature (°C)', '> 80', '> 60', '< 60')
    add_row('Battery (Voltage)', '< 2.5 V', '< 3.0 V', '> 3.0 V')
    
    doc.add_paragraph('\nNote: A "Critical" status overrides "Warning", and "Warning" overrides "Normal".')

    doc.add_heading('4.2. Command Validation Logic', level=2)
    doc.add_paragraph(
        'Responses from BLE devices are automatically graded:'
    )
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('PASS (P):').bold = True
    p.add_run(' Returned value is exactly "0".')

    p = doc.add_paragraph(style='List Bullet')
    p.add_run('FAIL (F):').bold = True
    p.add_run(' Returned value is negative (e.g., "-1").')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('UNKNOWN:').bold = True
    p.add_run(' Any other value or malformed string.')

    # --- 5. File Structure Overview ---
    doc.add_heading('5. File Structure & Purpose', level=1)
    doc.add_paragraph('A high-level guide to the project verification files:')

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'File Path'
    table.rows[0].cells[1].text = 'Description of Responsibility'
    
    files = [
        ("backend/app/main.py", "Server entry point. Orchestrates the startup sequence and routes requests."),
        ("backend/app/services/pcan_service.py", "Hardware Driver Layer. Contains the infinite loop thread for CAN reading."),
        ("backend/app/routers/*.py", "API Interface. Defines the HTTP endpoints exposed to the frontend."),
        ("frontend/src/pages/TPMSDashboard.jsx", "Primary UI Logic. Handles 3D rendering, chart data smoothing, and state updates."),
        ("frontend/src/services/api.js", "Network Layer. Centralized configuration for all HTTP/WebSocket calls."),
    ]
    
    for path, desc in files:
        row = table.add_row().cells
        row[0].text = path
        row[1].text = desc

    output_path = os.path.join(r"d:\Download (D Drive)\pcanlast (5)\tpms1", "Project_Detailed_Report.docx")
    doc.save(output_path)
    print(f"Document saved to: {output_path}")

if __name__ == "__main__":
    create_documentation()
